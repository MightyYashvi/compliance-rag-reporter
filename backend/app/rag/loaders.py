"""
Document loader for the Compliance RAG Reporter knowledge base.

Why metadata preservation matters:
    Every chunk produced downstream must carry enough provenance to let the
    report generator cite the exact source — document title, jurisdiction,
    page, and section. The citation/provenance layer has no separate lookup
    table; it relies entirely on the metadata attached here. This loader is
    the single place where raw file content is married to its provenance.
    Chunkers, the vector store, and the retrieval endpoint propagate that
    metadata unchanged.
"""
from __future__ import annotations

import csv
import hashlib
import logging
import re
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}

# Resolved at import time relative to this file's location so the default
# works whether the caller is at the repo root, inside backend/, or in tests.
_DEFAULT_KB_DIR = Path(__file__).resolve().parents[2] / "data" / "kb_raw"


# ---------------------------------------------------------------------------
# Document container
# ---------------------------------------------------------------------------

class KBDocument:
    """
    Thin wrapper around raw text + provenance metadata.

    Mirrors the LangChain Document interface (page_content / metadata) so
    the chunker can swap in langchain.schema.Document transparently later.
    """

    def __init__(self, page_content: str, metadata: dict[str, Any]) -> None:
        self.page_content = page_content
        self.metadata = metadata

    def __repr__(self) -> str:
        title = self.metadata.get("title", "untitled")
        return f"KBDocument(title={title!r}, chars={len(self.page_content)})"


# ---------------------------------------------------------------------------
# Metadata helpers
# ---------------------------------------------------------------------------

def _make_document_id(file_name: str) -> str:
    """Stable 16-char ID derived from filename — same file always gets same ID."""
    return hashlib.md5(file_name.encode()).hexdigest()[:16]


def _extract_version(lines: list[str]) -> str:
    """Look for Version/Revision/Edition markers in the first 10 lines."""
    patterns = [
        r"Version\s+([\w.]+)",
        r"Revision\s+([\w.]+)",
        r"([\d]{4})\s+Edition",
        r"Edition\s+([\d]{4})",
    ]
    for line in lines[:10]:
        for pat in patterns:
            m = re.search(pat, line, re.IGNORECASE)
            if m:
                return m.group(1)
    return ""


def _extract_standard_name(title: str) -> str:
    """Pull the uppercase acronym from parentheses, e.g. '... (SDIS)' -> 'SDIS'."""
    m = re.search(r"\(([A-Z]{2,})\)", title)
    return m.group(1) if m else title


def _infer_jurisdiction(file_name: str) -> str:
    if "singapore" in file_name.lower():
        return "Singapore"
    return "unknown"


def _build_metadata(path: Path, first_lines: list[str], page: int = 1) -> dict[str, Any]:
    """
    Construct a full provenance metadata dict from file path and content.

    Every field here becomes a citable attribute in the retrieval response.
    Adding a new field here automatically makes it available to the retriever
    and the report generator without touching any other layer.
    """
    title = next((ln.strip() for ln in first_lines if ln.strip()), path.stem)
    return {
        "document_id": _make_document_id(path.name),
        "title": title,
        "file_name": path.name,
        "file_path": str(path),
        "doc_type": "standard",
        "jurisdiction": _infer_jurisdiction(path.name),
        "standard_name": _extract_standard_name(title),
        "version": _extract_version(first_lines),
        "page": page,
        "source_type": path.suffix.lstrip(".").lower(),
    }


# ---------------------------------------------------------------------------
# Per-format loaders
# ---------------------------------------------------------------------------

def _load_txt(path: Path) -> list[KBDocument]:
    text = path.read_text(encoding="utf-8", errors="replace")
    meta = _build_metadata(path, text.splitlines())
    return [KBDocument(page_content=text, metadata=meta)]


def _load_md(path: Path) -> list[KBDocument]:
    return _load_txt(path)


def _load_pdf(path: Path) -> list[KBDocument]:
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.warning("pypdf not installed — skipping %s", path.name)
        return []

    reader = PdfReader(str(path))
    documents = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        meta = _build_metadata(path, text.splitlines(), page=page_num)
        documents.append(KBDocument(page_content=text, metadata=meta))
    return documents


def _load_docx(path: Path) -> list[KBDocument]:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.warning("python-docx not installed — skipping %s", path.name)
        return []

    doc = DocxDocument(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    meta = _build_metadata(path, text.splitlines())
    return [KBDocument(page_content=text, metadata=meta)]


_FORMAT_LOADERS = {
    ".txt": _load_txt,
    ".md": _load_md,
    ".pdf": _load_pdf,
    ".docx": _load_docx,
}


# ---------------------------------------------------------------------------
# Optional CSV metadata enrichment
# ---------------------------------------------------------------------------

def _read_metadata_csv(source_dir: Path) -> dict[str, dict[str, str]]:
    """
    Load metadata.csv from source_dir if present.

    CSV columns: file_name, title, doc_type, jurisdiction, standard_name, version
    Returns a dict keyed by file_name for O(1) lookup during enrichment.
    """
    csv_path = source_dir / "metadata.csv"
    if not csv_path.exists():
        return {}

    overrides: dict[str, dict[str, str]] = {}
    with csv_path.open(newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            fname = row.get("file_name", "").strip()
            if fname:
                overrides[fname] = {k: v.strip() for k, v in row.items() if v and v.strip()}
    logger.info("Loaded metadata overrides for %d files from metadata.csv", len(overrides))
    return overrides


def _apply_csv_overrides(
    metadata: dict[str, Any], overrides: dict[str, str]
) -> dict[str, Any]:
    """Overlay non-empty CSV values onto inferred metadata."""
    for key, value in overrides.items():
        if key != "file_name" and value:
            metadata[key] = value
    return metadata


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def load_kb_documents(
    source_dir: str | Path = _DEFAULT_KB_DIR,
) -> list[KBDocument]:
    """
    Load all supported documents from source_dir and return them as KBDocuments.

    Processing order:
      1. Read file content via the appropriate format loader.
      2. Infer provenance metadata from the filename and first few lines.
      3. Enrich metadata with any matching row from metadata.csv.
      4. Return documents sorted by filename for stable ordering.

    Unsupported files (.DS_Store, .csv, .xlsx, …) are skipped silently.
    Load errors for individual files are logged and skipped — one bad file
    will not abort the rest of the batch.
    """
    source_dir = Path(source_dir)
    if not source_dir.exists():
        logger.warning("KB directory not found: %s", source_dir)
        return []

    csv_overrides = _read_metadata_csv(source_dir)
    documents: list[KBDocument] = []

    for path in sorted(source_dir.iterdir()):
        if not path.is_file():
            continue

        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            logger.debug("Skipping unsupported file: %s", path.name)
            continue

        loader = _FORMAT_LOADERS.get(ext)
        if loader is None:
            continue

        try:
            docs = loader(path)
        except Exception:
            logger.exception("Failed to load %s — skipping", path.name)
            continue

        file_overrides = csv_overrides.get(path.name, {})
        if file_overrides:
            for doc in docs:
                doc.metadata = _apply_csv_overrides(doc.metadata, file_overrides)

        documents.extend(docs)
        logger.info("Loaded %d segment(s) from %s", len(docs), path.name)

    return documents


def summarize_loaded_documents(documents: list[KBDocument]) -> dict[str, Any]:
    """Return a lightweight summary suitable for logging and smoke-testing."""
    return {
        "document_count": len(documents),
        "files": [
            {
                "file_name": doc.metadata.get("file_name"),
                "title": doc.metadata.get("title"),
                "jurisdiction": doc.metadata.get("jurisdiction"),
                "doc_type": doc.metadata.get("doc_type"),
                "standard_name": doc.metadata.get("standard_name"),
                "version": doc.metadata.get("version"),
            }
            for doc in documents
        ],
    }
