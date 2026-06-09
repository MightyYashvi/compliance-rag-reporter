"""Report export  —  PERSON C deliverable.

Takes a structured report (sections + citations) and renders it to PDF or
DOCX. Pure functions returning bytes, so the Streamlit UI just offers a
download button.
"""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer


def _citation_text(citations: list[dict]) -> str:
    if not citations:
        return ""
    parts = [f"[{c['id']}] {c['source']} (chunk {c['chunk_id']})"
             for c in citations]
    return "  ".join(parts)


def to_docx(report: dict[str, Any]) -> bytes:
    """Render a structured report dict to DOCX bytes."""
    doc = Document()
    doc.add_heading(report.get("title", "Field Report"), level=0)

    meta = report.get("metadata", {})
    if meta:
        p = doc.add_paragraph()
        run = p.add_run(" | ".join(f"{k}: {v}" for k, v in meta.items()))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for section in report.get("sections", []):
        doc.add_heading(section["heading"], level=1)
        doc.add_paragraph(section["body"])
        cites = section.get("citations", [])
        if cites:
            cp = doc.add_paragraph()
            r = cp.add_run("Sources: " + _citation_text(cites))
            r.italic = True
            r.font.size = Pt(8)

    refs = report.get("references", [])
    if refs:
        doc.add_heading("References", level=1)
        for c in refs:
            doc.add_paragraph(
                f"[{c['id']}] {c['source']} — chunk {c['chunk_id']}: "
                f"\"{c.get('snippet', '')}\"")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf(report: dict[str, Any]) -> bytes:
    """Render a structured report dict to PDF bytes."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER,
                            topMargin=inch, bottomMargin=inch,
                            leftMargin=inch, rightMargin=inch)
    styles = getSampleStyleSheet()
    cite_style = ParagraphStyle("cite", parent=styles["Italic"],
                                fontSize=8, textColor="#555555")
    meta_style = ParagraphStyle("meta", parent=styles["Normal"],
                                fontSize=9, textColor="#666666")

    flow = [Paragraph(report.get("title", "Field Report"), styles["Title"])]

    meta = report.get("metadata", {})
    if meta:
        flow.append(Paragraph(" | ".join(f"{k}: {v}" for k, v in meta.items()),
                              meta_style))
    flow.append(Spacer(1, 12))

    for section in report.get("sections", []):
        flow.append(Paragraph(section["heading"], styles["Heading1"]))
        flow.append(Paragraph(section["body"], styles["BodyText"]))
        cites = section.get("citations", [])
        if cites:
            flow.append(Paragraph("Sources: " + _citation_text(cites),
                                  cite_style))
        flow.append(Spacer(1, 8))

    refs = report.get("references", [])
    if refs:
        flow.append(Paragraph("References", styles["Heading1"]))
        for c in refs:
            flow.append(Paragraph(
                f"[{c['id']}] {c['source']} — chunk {c['chunk_id']}: "
                f"\"{c.get('snippet', '')}\"", cite_style))

    doc.build(flow)
    return buf.getvalue()
